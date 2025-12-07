#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
症例要約をWord形式で出力するスクリプト

使用方法:
    python word_output_generator.py

必要なライブラリ:
    pip install python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import json
import sys
from pathlib import Path

def setup_japanese_font(doc):
    """日本語フォント（MS明朝）を設定"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'MS Mincho'
    font._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')
    font.size = Pt(11)
    return style

def create_case_summary_docx(case_data, output_path):
    """
    症例要約をWord形式で出力
    
    Args:
        case_data: 症例データ（辞書形式）
        output_path: 出力ファイルパス
    """
    doc = Document()
    
    # ページ設定（A4）
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4 height
    section.page_width = Inches(8.27)     # A4 width
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    
    # フォント設定
    style = setup_japanese_font(doc)
    
    # 基本情報セクション
    p = doc.add_paragraph()
    run = p.add_run(f"症例番号: {case_data.get('case_number', '')}")
    run.bold = True
    run.font.size = Pt(11)
    
    if case_data.get('is_designated_disease'):
        p.add_run(" ○")  # 指定疾患マーク
    
    p = doc.add_paragraph()
    p.add_run(f"分野番号: {case_data.get('field_number', '')}")
    
    p = doc.add_paragraph()
    p.add_run(f"患者ID: {case_data.get('patient_id', '')}")
    
    p = doc.add_paragraph()
    inpatient_text = '入院症例' if case_data.get('is_inpatient', True) else '外来症例'
    p.add_run(f"入院・外来: {inpatient_text}")
    
    p = doc.add_paragraph()
    p.add_run(f"受け持ち期間: {case_data.get('care_period', '')}")
    
    p = doc.add_paragraph()
    p.add_run(f"年齢: {case_data.get('age', '')}")
    
    p = doc.add_paragraph()
    gender_text = '男' if case_data.get('gender') == 'male' else '女'
    p.add_run(f"性別: {gender_text}")
    
    p = doc.add_paragraph()
    p.add_run(f"転帰: {case_data.get('outcome', '')}")
    
    # 本文セクション（30行以内に収める）
    doc.add_paragraph()  # 空行
    
    # 主訴
    if case_data.get('chief_complaint'):
        p = doc.add_paragraph()
        run = p.add_run("【主訴】")
        run.bold = True
        p.add_run(f" {case_data['chief_complaint']}")
    
    # 現病歴
    if case_data.get('present_illness'):
        p = doc.add_paragraph()
        run = p.add_run("【現病歴】")
        run.bold = True
        p.add_run(f" {case_data['present_illness']}")
    
    # 診察所見
    exam_label = "【入院時診察所見】" if case_data.get('is_inpatient', True) else "【来院時診察所見】"
    if case_data.get('physical_examination'):
        p = doc.add_paragraph()
        run = p.add_run(exam_label)
        run.bold = True
        p.add_run(f" {case_data['physical_examination']}")
    
    # 検査所見
    lab_label = "【入院時検査所見】" if case_data.get('is_inpatient', True) else "【来院時検査所見】"
    if case_data.get('laboratory_findings'):
        p = doc.add_paragraph()
        run = p.add_run(lab_label)
        run.bold = True
        p.add_run(f" {case_data['laboratory_findings']}")
    
    # 鑑別診断（必須）
    if case_data.get('differential_diagnoses'):
        p = doc.add_paragraph()
        run = p.add_run("【鑑別診断】")
        run.bold = True
        for i, dd in enumerate(case_data['differential_diagnoses'], 1):
            if i > 1:
                p.add_run(" ")
            p.add_run(f"{i}. {dd}。")
    
    # 症例の問題点（推奨）
    if case_data.get('problem_points'):
        p = doc.add_paragraph()
        run = p.add_run("【症例の問題点】")
        run.bold = True
        p.add_run(f" {case_data['problem_points']}")
    
    # 入院後経過
    course_label = "【入院後経過】" if case_data.get('is_inpatient', True) else "【来院後経過】"
    if case_data.get('hospital_course'):
        p = doc.add_paragraph()
        run = p.add_run(course_label)
        run.bold = True
        p.add_run(f" {case_data['hospital_course']}")
    
    # 家族への説明（必須）
    if case_data.get('family_explanation'):
        p = doc.add_paragraph()
        run = p.add_run("【家族への説明・指示】")
        run.bold = True
        p.add_run(f" {case_data['family_explanation']}")
    
    # 退院後の経過（推奨）
    if case_data.get('post_discharge_course'):
        p = doc.add_paragraph()
        run = p.add_run("【退院後の経過】")
        run.bold = True
        p.add_run(f" {case_data['post_discharge_course']}")
    
    # 保存
    doc.save(output_path)
    print(f"✅ Wordファイルを出力しました: {output_path}")

def create_all_case_summaries(cases_data, output_path):
    """
    複数症例を1つのWordファイルに出力
    
    Args:
        cases_data: 症例データのリスト
        output_path: 出力ファイルパス
    """
    doc = Document()
    
    # ページ設定
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    
    # フォント設定
    style = setup_japanese_font(doc)
    
    # 症例を分野番号・症例番号順にソート
    sorted_cases = sorted(cases_data, key=lambda x: (
        x.get('field_number', 0),
        x.get('case_number', 0)
    ))
    
    # 各症例を追加
    for i, case in enumerate(sorted_cases):
        if i > 0:
            # ページ区切り（最後の症例以外）
            doc.add_page_break()
        
        # 症例区切り線
        p = doc.add_paragraph("=" * 50)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 症例要約を追加
        add_case_summary_section(doc, case)
    
    doc.save(output_path)
    print(f"✅ 全{len(cases_data)}症例を出力しました: {output_path}")

def add_case_summary_section(doc, case_data):
    """症例要約セクションを追加"""
    # 基本情報
    p = doc.add_paragraph()
    run = p.add_run(f"症例番号: {case_data.get('case_number', '')}")
    run.bold = True
    
    if case_data.get('is_designated_disease'):
        p.add_run(" ○")
    
    p = doc.add_paragraph()
    p.add_run(f"分野番号: {case_data.get('field_number', '')}")
    
    p = doc.add_paragraph()
    p.add_run(f"患者ID: {case_data.get('patient_id', '')}")
    
    p = doc.add_paragraph()
    inpatient_text = '入院症例' if case_data.get('is_inpatient', True) else '外来症例'
    p.add_run(f"入院・外来: {inpatient_text}")
    
    p = doc.add_paragraph()
    p.add_run(f"受け持ち期間: {case_data.get('care_period', '')}")
    
    p = doc.add_paragraph()
    p.add_run(f"年齢: {case_data.get('age', '')}")
    
    p = doc.add_paragraph()
    gender_text = '男' if case_data.get('gender') == 'male' else '女'
    p.add_run(f"性別: {gender_text}")
    
    p = doc.add_paragraph()
    p.add_run(f"転帰: {case_data.get('outcome', '')}")
    
    doc.add_paragraph()  # 空行
    
    # 本文
    if case_data.get('chief_complaint'):
        p = doc.add_paragraph()
        run = p.add_run("【主訴】")
        run.bold = True
        p.add_run(f" {case_data['chief_complaint']}")
    
    if case_data.get('present_illness'):
        p = doc.add_paragraph()
        run = p.add_run("【現病歴】")
        run.bold = True
        p.add_run(f" {case_data['present_illness']}")
    
    exam_label = "【入院時診察所見】" if case_data.get('is_inpatient', True) else "【来院時診察所見】"
    if case_data.get('physical_examination'):
        p = doc.add_paragraph()
        run = p.add_run(exam_label)
        run.bold = True
        p.add_run(f" {case_data['physical_examination']}")
    
    lab_label = "【入院時検査所見】" if case_data.get('is_inpatient', True) else "【来院時検査所見】"
    if case_data.get('laboratory_findings'):
        p = doc.add_paragraph()
        run = p.add_run(lab_label)
        run.bold = True
        p.add_run(f" {case_data['laboratory_findings']}")
    
    if case_data.get('differential_diagnoses'):
        p = doc.add_paragraph()
        run = p.add_run("【鑑別診断】")
        run.bold = True
        for i, dd in enumerate(case_data['differential_diagnoses'], 1):
            if i > 1:
                p.add_run(" ")
            p.add_run(f"{i}. {dd}。")
    
    if case_data.get('problem_points'):
        p = doc.add_paragraph()
        run = p.add_run("【症例の問題点】")
        run.bold = True
        p.add_run(f" {case_data['problem_points']}")
    
    course_label = "【入院後経過】" if case_data.get('is_inpatient', True) else "【来院後経過】"
    if case_data.get('hospital_course'):
        p = doc.add_paragraph()
        run = p.add_run(course_label)
        run.bold = True
        p.add_run(f" {case_data['hospital_course']}")
    
    if case_data.get('family_explanation'):
        p = doc.add_paragraph()
        run = p.add_run("【家族への説明・指示】")
        run.bold = True
        p.add_run(f" {case_data['family_explanation']}")
    
    if case_data.get('post_discharge_course'):
        p = doc.add_paragraph()
        run = p.add_run("【退院後の経過】")
        run.bold = True
        p.add_run(f" {case_data['post_discharge_course']}")

def main():
    """メイン関数"""
    if len(sys.argv) < 2:
        print("使用方法:")
        print("  python word_output_generator.py <JSONファイル>")
        print("  python word_output_generator.py <JSONファイル> --all")
        print("\n例:")
        print("  python word_output_generator.py case1.json")
        print("  python word_output_generator.py cases.json --all")
        sys.exit(1)
    
    json_path = Path(sys.argv[1])
    if not json_path.exists():
        print(f"❌ ファイルが見つかりません: {json_path}")
        sys.exit(1)
    
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    if '--all' in sys.argv or isinstance(data, list):
        # 複数症例を一括出力
        if isinstance(data, list):
            cases_data = data
        else:
            cases_data = [data]
        
        output_path = json_path.stem + '_all.docx'
        create_all_case_summaries(cases_data, output_path)
    else:
        # 単一症例を出力
        output_path = json_path.stem + '.docx'
        create_case_summary_docx(data, output_path)

if __name__ == '__main__':
    main()
